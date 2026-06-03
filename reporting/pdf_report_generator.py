from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── helpers ──────────────────────────────────────────────────────────────────

def _env_int(key: str, default: int = 0) -> int:
    try:
        return int(os.getenv(key, default))
    except (ValueError, TypeError):
        return default


def _set_cell_text(cell, text: str, bold: bool = False, font_size: int = 10) -> None:
    """Escribe texto simple en una celda (una sola línea)."""
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)


def _set_cell_lines(cell, lines: list[str], bold_first: bool = False, font_size: int = 10) -> None:
    """Escribe varias líneas en una celda, cada una como párrafo independiente."""
    cell.text = ""
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        run = p.add_run(line)
        run.bold = bold_first and i == 0
        run.font.size = Pt(font_size)


def _decision_checkbox(label: str, checked: bool) -> str:
    return f"{'☑' if checked else '☐'}  {label}"


def _compute_certification(summary: dict) -> dict:
    """Reutiliza la misma lógica de _status_data sin importar el módulo completo."""
    total    = int(summary.get("total",    0) or 0)
    failures = int(summary.get("failures", 0) or 0)
    errors   = int(summary.get("errors",   0) or 0)
    skipped  = int(summary.get("skipped",  0) or 0)

    if total > 0 and failures == 0 and errors == 0 and skipped == 0:
        return {"certification": "CERTIFICABLE", "limitado": False}
    if total == 0:
        return {"certification": "NO CERTIFICABLE", "limitado": False}
    if skipped > 0 and failures == 0 and errors == 0:
        return {"certification": "CERTIFICABLE CON ALCANCE LIMITADO", "limitado": True}
    return {"certification": "NO CERTIFICABLE", "limitado": False}


# ── builder principal ─────────────────────────────────────────────────────────

def generate_pdf_report(
    summary: dict,
    feature_meta: dict,
    narrative: dict,
    output_dir: Path,
) -> Path | None:
    """
    Genera reporte-cierre.docx a partir del summary + metadata del feature,
    luego lo convierte a PDF con docx2pdf (requiere MS Word en Windows).

    Retorna la ruta al PDF generado, o None si falla la conversión.
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ── estilo base ───────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ══════════════════════════════════════════════════════════════════════════
    # TÍTULO
    # ══════════════════════════════════════════════════════════════════════════
    title_p = doc.add_heading("Reporte de Cierre de Calidad", level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        "Documento formal de cierre del ciclo de pruebas, resultados finales "
        "y recomendación de liberación."
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. IDENTIFICACIÓN DEL RELEASE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Identificación del Release", level=2)

    started_at = summary.get("started_at", "")
    finished_at = summary.get("finished_at", datetime.now().strftime("%d/%m/%Y"))
    if started_at:
        periodo = f"{started_at} - {finished_at}"
    else:
        periodo = finished_at

    id_rows = [
        ("Producto / Proyecto", feature_meta.get("producto", "")),
        ("Vertical",            feature_meta.get("vertical", "")),
        ("Release / Versión",   feature_meta.get("release", "")),
        ("Sprint(s) cubiertos", feature_meta.get("jira", "")),
        ("Período de pruebas",  periodo),
        ("QA Lead responsable", feature_meta.get("qa lead", "")),
        ("Product Owner",       feature_meta.get("product owner", "")),
        ("Fecha de cierre",     datetime.now().strftime("%d / %m / %Y")),
    ]

    table = doc.add_table(rows=len(id_rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(id_rows):
        _set_cell_text(table.cell(i, 0), label, bold=True)
        _set_cell_text(table.cell(i, 1), value)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. RESUMEN EJECUTIVO
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Resumen Ejecutivo", level=2)
    exec_summary = narrative.get("executive_summary", "")
    doc.add_paragraph(exec_summary)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. ALCANCE PROBADO
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Alcance Probado", level=2)

    test_cases = summary.get("test_cases", [])
    alcance_table = doc.add_table(rows=1, cols=3)
    alcance_table.style = "Table Grid"

    headers = ["Historia / Feature", "Tipo de prueba", "Resultado"]
    for i, h in enumerate(headers):
        _set_cell_text(alcance_table.cell(0, i), h, bold=True)

    feature_name = feature_meta.get(
        "producto",
        summary.get("execution_name", "Automatización QA"),
    )
    jira_link = feature_meta.get("jira", "")

    seen: set[str] = set()
    for tc in test_cases:
        status = (tc.get("status") or "").upper()
        result = "OK" if status == "PASSED" else ("FAIL" if status == "FAILED" else status)
        key = f"{jira_link}|Functional|{result}"
        if key in seen:
            continue
        seen.add(key)
        row = alcance_table.add_row()
        _set_cell_text(row.cells[0], jira_link or feature_name)
        _set_cell_text(row.cells[1], "Functional")
        _set_cell_text(row.cells[2], result)

    if not test_cases:
        row = alcance_table.add_row()
        _set_cell_text(row.cells[0], jira_link or feature_name)
        _set_cell_text(row.cells[1], "Functional")
        _set_cell_text(row.cells[2], "N/A")

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. RESULTADOS FINALES DE EJECUCIÓN
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Resultados Finales de Ejecución", level=2)

    total    = int(summary.get("total",    0) or 0)
    passed   = int(summary.get("passed",   0) or 0)
    failures = int(summary.get("failures", 0) or 0)
    errors   = int(summary.get("errors",   0) or 0)
    skipped  = int(summary.get("skipped",  0) or 0)
    failed   = failures + errors

    pct_ok = f"{round(passed / total * 100)}%" if total else "0%"

    result_headers = ["Indicador", "Planificado", "Ejecutado", "Exitoso", "Fallido", "% Éxito"]
    result_rows = [
        ("Casos de prueba",       total, total, passed, failed, pct_ok),
        ("Historias de usuario",  1,     1,     1 if failed == 0 else 0, 1 if failed > 0 else 0, "100%" if failed == 0 else "0%"),
        ("Pruebas de regresión",  total, total, passed, failed, pct_ok),
        ("Pruebas de integración","",    "",    "",     "",     ""),
    ]

    res_table = doc.add_table(rows=1, cols=6)
    res_table.style = "Table Grid"
    for i, h in enumerate(result_headers):
        _set_cell_text(res_table.cell(0, i), h, bold=True)

    for row_data in result_rows:
        row = res_table.add_row()
        for i, val in enumerate(row_data):
            _set_cell_text(row.cells[i], str(val))

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. COBERTURA DE PRUEBAS
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Cobertura de Pruebas", level=2)

    cobertura_table = doc.add_table(rows=4, cols=2)
    cobertura_table.style = "Table Grid"

    automated = sum(1 for tc in test_cases)
    cob_rows = [
        ("Cobertura funcional",                    pct_ok),
        ("Cobertura de regresión",                 pct_ok),
        ("Cobertura por criterios de aceptación",  pct_ok),
        ("Casos automatizados vs manuales",         f"{automated} automatizados / 0 manuales"),
    ]
    for i, (label, value) in enumerate(cob_rows):
        _set_cell_text(cobertura_table.cell(i, 0), label, bold=True)
        _set_cell_text(cobertura_table.cell(i, 1), value)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. ESTADO FINAL DE DEFECTOS
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Estado Final de Defectos", level=2)

    det_crit  = _env_int("DEFECTOS_CRITICOS")
    det_alto  = _env_int("DEFECTOS_ALTOS")
    det_medio = _env_int("DEFECTOS_MEDIOS")
    det_bajo  = _env_int("DEFECTOS_BAJOS")

    cer_crit  = _env_int("DEFECTOS_CRITICOS_CERRADOS")
    cer_alto  = _env_int("DEFECTOS_ALTOS_CERRADOS")
    cer_medio = _env_int("DEFECTOS_MEDIOS_CERRADOS")
    cer_bajo  = _env_int("DEFECTOS_BAJOS_CERRADOS")

    dif_crit  = _env_int("DEFECTOS_CRITICOS_DIFERIDOS")
    dif_alto  = _env_int("DEFECTOS_ALTOS_DIFERIDOS")
    dif_medio = _env_int("DEFECTOS_MEDIOS_DIFERIDOS")
    dif_bajo  = _env_int("DEFECTOS_BAJOS_DIFERIDOS")

    justif = os.getenv("DEFECTOS_DIFERIDOS_JUSTIFICACION", "")

    def abiertos(det: int, cer: int, dif: int) -> int:
        return max(0, det - cer - dif)

    def row_data(label: str, det: int, cer: int, dif: int) -> tuple:
        ab = abiertos(det, cer, dif)
        return (label, det, cer, ab, dif, justif if dif > 0 else "")

    total_det = det_crit + det_alto + det_medio + det_bajo
    total_cer = cer_crit + cer_alto + cer_medio + cer_bajo
    total_dif = dif_crit + dif_alto + dif_medio + dif_bajo
    total_ab  = abiertos(total_det, total_cer, total_dif)

    def_headers = ["Severidad", "Detectados", "Cerrados", "Abiertos", "Diferidos", "Justificación de diferidos"]
    def_rows = [
        row_data("Crítica", det_crit,  cer_crit,  dif_crit),
        row_data("Alta",    det_alto,  cer_alto,  dif_alto),
        row_data("Media",   det_medio, cer_medio, dif_medio),
        row_data("Baja",    det_bajo,  cer_bajo,  dif_bajo),
        ("TOTAL", total_det, total_cer, total_ab, total_dif, ""),
    ]

    def_table = doc.add_table(rows=1, cols=6)
    def_table.style = "Table Grid"
    for i, h in enumerate(def_headers):
        _set_cell_text(def_table.cell(0, i), h, bold=True)

    for row_d in def_rows:
        row = def_table.add_row()
        for i, val in enumerate(row_d):
            bold = row_d[0] == "TOTAL"
            _set_cell_text(row.cells[i], str(val), bold=bold)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. DEUDA DE CALIDAD GENERADA
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Deuda de Calidad Generada", level=2)

    closing = narrative.get("closing_observation", "")
    if closing:
        for line in closing.strip().splitlines():
            line = line.strip().lstrip("-•·").strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("Sin deuda de calidad identificada.")

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 8. RECOMENDACIÓN DE CALIDAD
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Recomendación de Calidad", level=2)

    cert_data  = _compute_certification(summary)
    cert_str   = cert_data["certification"]
    is_no_apt  = "NO CERTIFICABLE" in cert_str
    is_limitado = cert_data["limitado"] or total_ab > 0
    is_apto    = not is_no_apt and not is_limitado

    rec_table = doc.add_table(rows=1, cols=1)
    rec_table.style = "Table Grid"
    cell = rec_table.cell(0, 0)

    checkboxes = [
        ("APTO PARA LIBERACIÓN",                          is_apto),
        ("APTO CON OBSERVACIONES  (ver sección de deuda)", is_limitado),
        ("NO APTO PARA LIBERACIÓN",                       is_no_apt),
    ]

    for i, (label, checked) in enumerate(checkboxes):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(_decision_checkbox(label, checked))
        run.bold = checked
        run.font.size = Pt(10)

    justif_text = narrative.get("executive_conclusion", "").strip()
    p_just = cell.add_paragraph()
    r1 = p_just.add_run("Justificación: ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p_just.add_run(justif_text)
    r2.font.size = Pt(10)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 9. FIRMAS DE CIERRE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Firmas de Cierre", level=2)
    doc.add_paragraph(
        "Doble gate de calidad: el Product Owner aprueba el release desde la "
        "perspectiva de producto y QA formaliza la liberación."
    )

    qa_lead   = feature_meta.get("qa lead", "QA Responsable")
    po        = feature_meta.get("product owner", "Product Owner")
    gerencia  = "Gerencia de Producto & QA"

    # Tabla QA + PO
    firma_table = doc.add_table(rows=2, cols=2)
    firma_table.style = "Table Grid"

    _set_cell_lines(firma_table.cell(0, 0), [""], bold_first=False)   # espacio para firma
    _set_cell_lines(firma_table.cell(0, 1), [""], bold_first=False)

    _set_cell_lines(
        firma_table.cell(1, 0),
        [qa_lead, "QA Responsable", "Cierre técnico de calidad"],
        bold_first=True,
    )
    _set_cell_lines(
        firma_table.cell(1, 1),
        [po, "Product Owner", "Aprobación de producto"],
        bold_first=True,
    )

    doc.add_paragraph()

    # Tabla Gerencia
    gerencia_table = doc.add_table(rows=2, cols=1)
    gerencia_table.style = "Table Grid"
    _set_cell_lines(gerencia_table.cell(0, 0), [""], bold_first=False)  # espacio para firma
    _set_cell_lines(
        gerencia_table.cell(1, 0),
        [gerencia, "Visto bueno"],
        bold_first=True,
    )

    # ── Guardar .docx ─────────────────────────────────────────────────────────
    docx_path = output_dir / "reporte-cierre.docx"
    doc.save(docx_path)
    print(f"\n[PDF Report] DOCX generado: {docx_path}")

    # ── Convertir a PDF ───────────────────────────────────────────────────────
    pdf_path = output_dir / "reporte-cierre.pdf"
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        print(f"[PDF Report] PDF generado: {pdf_path}")
        return pdf_path
    except ImportError:
        print("[PDF Report] docx2pdf no instalado. Solo se generó el DOCX.")
        print("             Instala con: pip install docx2pdf")
        return None
    except Exception as exc:
        print(f"[PDF Report] Error al convertir a PDF: {exc}")
        print(f"             El DOCX está disponible en: {docx_path}")
        return None
